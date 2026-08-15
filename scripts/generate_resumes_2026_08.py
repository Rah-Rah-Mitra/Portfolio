from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
GENERATED = ROOT / "public" / "resume" / "generated"
EDITION_FROM = "2026-07"
EDITION_TO = "2026-08"

CONTACT_LINKS = [
    ("Singapore", None),
    ("mitrarahul2002@gmail.com", "mailto:mitrarahul2002@gmail.com"),
    ("linkedin.com/in/rahulmitra-dev", "https://www.linkedin.com/in/rahulmitra-dev"),
    ("github.com/Rah-Rah-Mitra", "https://github.com/Rah-Rah-Mitra"),
    ("rahul-mitra.com", "https://rahul-mitra.com/"),
]

ABBOTT_DIGITAL_TWIN_SOURCE = (
    "- Built a Python digital twin simulation and constraint-programming optimizer for an NP-hard hybrid flow shop scheduling problem, "
    "translating production complexity into decision-ready schedules."
)
ABBOTT_DIGITAL_TWIN_GENERAL_SOURCE = (
    "- Built a Python digital-twin simulation and constraint-programming optimizer for an NP-hard hybrid flow-shop scheduling problem, "
    "translating production complexity into decision-ready schedules."
)
ABBOTT_APC_SOURCE = (
    "- Architected an Azure-based web simulator for Advanced Process Control systems across global manufacturing plants, "
    "using Docker containerization and Bitbucket CI/CD for repeatable releases."
)
GLOBAL_INLINE_REPLACEMENTS = {
    "Abbott Laboratories Singapore Pte Ltd": "Abbott Laboratories",
}
INLINE_REPLACEMENTS: dict[str, dict[str, str]] = {
    "civic-tech-solution-architect": {"Azure APC Web Simulator": "APC Simulator Cloud Operations"},
    "solution-architect": {"Azure APC Web Simulator": "APC Simulator Cloud Operations"},
}

REPLACEMENTS: dict[str, dict[str, str]] = {
    "ai-engineer": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling, translating production objectives and constraints into decision-ready schedules.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "Skills: Python; BERT; DNN; RAG; AI Agents; OpenAI GPT-4; Gemini; Vector Search; SEALION; ASPIRE 2A; LangChain; Computer Vision; Multi-View Geometry; Structure-from-Motion; Model Evaluation.":
        "Skills: Python; BERT; DNN; RAG; AI Agents; GPT-4; Gemini; Vector Search; SEALION; ASPIRE 2A; 3D Computer Vision; Projective & Epipolar Geometry; Absolute Pose; Structure-from-Motion; Bundle Adjustment; Multi-View Stereo; Model Evaluation.",
    },
    "civic-tech-solution-architect": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling, translating production objectives and constraints into decision-ready schedules.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "- Designed a cloud-hosted process-control simulation surface with Azure deployment, Docker packaging, CI/CD pipelines, and manufacturing stakeholder feedback loops.":
        "- Owned cloud hosting and production support for a team-built process-control simulator, Docker-packaged on Azure App Service for active internal manufacturing and engineer-training use.",
    },
    "cyber-security": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline with staged audit outputs to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        ABBOTT_APC_SOURCE:
        "- Productionized and supported a team-built APC simulator for live internal users on Azure App Service, applying access restrictions, runtime/session hardening, and Docker packaging.",
    },
    "general": {
        ABBOTT_DIGITAL_TWIN_GENERAL_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling, translating production objectives and constraints into decision-ready schedules.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "Operations Research / Optimization: Constraint Programming, Scheduling Optimization, Hybrid Flow Shop, Digital Twin Simulation, Simulation, Dijkstra / SSSP / APSP, Decision Analytics, Statistics, Jupyter, Mathematical Modeling (Linear Algebra, Probability, Calculus)":
        "Operations Research / Optimization: CP-SAT, Hybrid Flow-Shop Scheduling, SimPy, Digital Twin Simulation, Robust Optimization Research, Graph Optimization, Network Flow, Dijkstra / SSSP / APSP, Decision Analytics, Statistics, Mathematical Modeling",
        "AI/ML: BERT, DNN, Transformers, RAG / RAG Architecture, AI Agents, OpenAI GPT-4, Gemini, SEALION, LangChain, Vector Search, Computer Vision, Multi-View Geometry, Structure-from-Motion, Bundle Adjustment, Model Evaluation, NLP, Deep Reinforcement Learning (PPO, A2C, DDPG, DQN, Double DQN), ASPIRE 2A / HPC / CUDA / distributed training":
        "AI/ML: BERT, DNN, Transformers, RAG, AI Agents, GPT-4, Gemini, SEALION, LangChain, Vector Search, 3D Computer Vision, Projective & Epipolar Geometry, Absolute Pose, Structure-from-Motion, Bundle Adjustment, Multi-View Stereo, Model Evaluation, NLP, Deep RL (PPO, A2C, DDPG, DQN), ASPIRE 2A / HPC / CUDA / distributed training",
    },
    "software-engineer": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling, translating production objectives and constraints into decision-ready schedules.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "Skills: Python; TypeScript; React; Next.js; Three.js; FastAPI; Flask; aiohttp; asyncio; Java; Docker; CI/CD; Azure; Singpass/Myinfo; Codex.":
        "Skills: Python; TypeScript; React; Next.js; Three.js; FastAPI; aiohttp/asyncio; Docker; CI/CD; Azure; 3D CV: projective & epipolar geometry, absolute pose, SfM, bundle adjustment, multi-view stereo.",
    },
    "solution-architect": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling, translating production objectives and constraints into decision-ready schedules.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "- Designed a cloud-hosted process-control simulation surface with Azure deployment, Docker packaging, CI/CD pipelines, and manufacturing stakeholder feedback loops.":
        "- Productionized a team-built process-control simulator for active internal use through Azure App Service, Docker packaging, controlled access, runtime support, and deployment handover.",
    },
    "operations-research-engineer": {
        ABBOTT_DIGITAL_TWIN_SOURCE:
        "- Developed a SimPy discrete-event digital twin and OR-Tools CP-SAT optimizer for hybrid flow-shop scheduling; compared heuristic, MIP, and genetic-algorithm approaches and researched robust optimization.",
        ABBOTT_APC_SOURCE:
        "- Engineered a non-destructive, 15-stage changeover-data pipeline to normalize operator- and machine-generated anomalies, validating it against five years of previously unseen raw data with zero execution failures.",
        "- Modeled production scheduling as a digital twin simulation and constraint-programming optimization problem, connecting ISE methods with deployment-oriented Python tooling.":
        "- Modeled hybrid flow-shop scheduling as a SimPy digital twin with explicit objectives and constraints, using CP-SAT, simulation, and deployment-oriented Python tooling.",
        "Skills: Constraint Programming; Scheduling Optimization; Hybrid Flow Shop; Digital Twin Simulation; Dijkstra; SSSP/APSP; Python; Jupyter; Statistics; Simulation; Excel VBA; Decision Analytics.":
        "Skills: CP-SAT; Hybrid Flow-Shop Scheduling; SimPy; Robust Optimization Research; Digital Twin Simulation; Objective Functions & Constraints; Graph Optimization; Network Flow; Dijkstra; Python; Statistics; Excel VBA; Decision Analytics.",
    },
}

ADDITIONAL_ABBOTT_BULLETS: dict[str, tuple[str, ...]] = {
    "ai-engineer": (
        "- Productionized and supported a team-built APC simulator for live manufacturing and engineer training by packaging it in Docker and deploying it to Azure App Service.",
        "- Delivered practical AI upskilling to the regional engineering workforce for day-to-day operational use.",
    ),
    "civic-tech-solution-architect": (
        "- Productionized and supported a team-built APC simulator for live manufacturing and engineer training by packaging it in Docker and deploying it to Azure App Service.",
        "- Delivered practical AI upskilling to the regional engineering workforce for day-to-day operational use.",
    ),
    "general": (
        "- Productionized and supported a team-built APC simulator for live manufacturing and engineer training by packaging it in Docker and deploying it to Azure App Service.",
        "- Delivered practical AI upskilling to the regional engineering workforce for day-to-day operational use.",
    ),
    "software-engineer": (
        "- Productionized and supported a team-built APC simulator for live manufacturing and engineer training, resolving runtime issues, packaging it in Docker, and deploying it to Azure App Service.",
        "- Delivered practical AI upskilling to the regional engineering workforce for day-to-day operational use.",
    ),
    "solution-architect": (
        "- Owned cloud deployment and production support for a team-built APC simulator used in manufacturing and engineer training, with Docker packaging, Azure App Service hosting, and documented handover.",
        "- Delivered practical AI upskilling to the regional engineering workforce for day-to-day operational use.",
    ),
    "operations-research-engineer": (
        "- Productionized and supported a team-built APC simulator for live manufacturing and engineer training by packaging it in Docker and deploying it to Azure App Service.",
    ),
}


def copy_run_properties(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_hyperlink(paragraph, label: str, target: str, sample_run) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = deepcopy(sample_run._r.rPr) if sample_run is not None and sample_run._r.rPr is not None else OxmlElement("w:rPr")
    color = properties.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        properties.append(color)
    color.set(qn("w:val"), "0B514C")
    underline = properties.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        properties.append(underline)
    underline.set(qn("w:val"), "single")
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_contact_line(paragraph) -> None:
    sample_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph_content(paragraph)
    for index, (label, target) in enumerate(CONTACT_LINKS):
        if index:
            separator = paragraph.add_run(" | ")
            copy_run_properties(sample_run, separator)
        if target:
            add_hyperlink(paragraph, label, target, sample_run)
        else:
            run = paragraph.add_run(label)
            copy_run_properties(sample_run, run)


def replace_paragraph_text(paragraph, text: str) -> None:
    sample_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    copy_run_properties(sample_run, run)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    cloned_element = deepcopy(paragraph._p)
    paragraph._p.addnext(cloned_element)
    inserted = Paragraph(cloned_element, paragraph._parent)
    replace_paragraph_text(inserted, text)
    return inserted


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def generate(slug: str) -> Path:
    source = GENERATED / f"rahul-mitra-{slug}-{EDITION_FROM}.docx"
    target = GENERATED / f"rahul-mitra-{slug}-{EDITION_TO}.docx"
    document = Document(source)
    replacements = REPLACEMENTS.get(slug, {})
    found_contact = False
    found_replacements: set[str] = set()
    found_inline_replacements: set[str] = set()
    abbott_insert_anchor = None
    inline_replacements = {**GLOBAL_INLINE_REPLACEMENTS, **INLINE_REPLACEMENTS.get(slug, {})}

    paragraphs = list(iter_paragraphs(document))
    for paragraph in paragraphs:
        original = paragraph.text.strip()
        if "mitrarahul2002@gmail.com" in original and ("vercel.app" in original or "rahul-mitra.com" in original):
            replace_contact_line(paragraph)
            found_contact = True
            continue
        if original in replacements:
            replace_paragraph_text(paragraph, replacements[original])
            found_replacements.add(original)
            if original == ABBOTT_APC_SOURCE:
                abbott_insert_anchor = paragraph

        for run in paragraph.runs:
            for source_text, target_text in inline_replacements.items():
                if source_text in run.text:
                    run.text = run.text.replace(source_text, target_text)
                    found_inline_replacements.add(source_text)

    if not found_contact:
        raise RuntimeError(f"Contact line not found in {source.name}")
    missing = set(replacements) - found_replacements
    if missing:
        raise RuntimeError(f"Expected content was not found in {source.name}: {sorted(missing)}")
    missing_inline = set(inline_replacements) - found_inline_replacements
    if missing_inline:
        raise RuntimeError(f"Expected inline content was not found in {source.name}: {sorted(missing_inline)}")

    additional_bullets = ADDITIONAL_ABBOTT_BULLETS.get(slug, ())
    if additional_bullets:
        if abbott_insert_anchor is None:
            raise RuntimeError(f"Abbott insertion anchor not found in {source.name}")
        for bullet in additional_bullets:
            abbott_insert_anchor = insert_paragraph_after(abbott_insert_anchor, bullet)

    paragraphs = list(iter_paragraphs(document))
    for current, following in zip(paragraphs, paragraphs[1:]):
        current_text = current.text.strip()
        if current_text and not current_text.startswith("- ") and following.text.strip().startswith("- "):
            current.paragraph_format.keep_with_next = True

    document.core_properties.author = "Rahul Mitra"
    document.core_properties.last_modified_by = "Rahul Mitra"
    document.core_properties.subject = f"Rahul Mitra — {slug.replace('-', ' ').title()} résumé"
    document.core_properties.keywords = "Rahul Mitra, résumé, rahul-mitra.com"
    document.save(target)
    return target


def main() -> None:
    sources = sorted(GENERATED.glob(f"rahul-mitra-*-{EDITION_FROM}.docx"))
    if len(sources) != 7:
        raise RuntimeError(f"Expected seven source résumés, found {len(sources)}")
    generated = [generate(source.name.removeprefix("rahul-mitra-").removesuffix(f"-{EDITION_FROM}.docx")) for source in sources]
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
